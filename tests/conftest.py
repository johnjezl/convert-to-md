"""Shared test fixtures — generates sample .docx and .pdf files programmatically."""

from __future__ import annotations

from pathlib import Path

import fitz  # pymupdf
import pytest
from docx import Document
from docx.shared import Inches


@pytest.fixture
def tmp_output(tmp_path: Path) -> Path:
    """Return a temp directory for conversion output."""
    return tmp_path


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Create a simple .docx with headings, paragraph, bold, and a list."""
    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("This is a test paragraph with some text.")
    p = doc.add_paragraph()
    run = p.add_run("Bold text")
    run.bold = True
    p.add_run(" and normal text.")
    doc.add_heading("Section Two", level=2)
    doc.add_paragraph("Item one", style="List Bullet")
    doc.add_paragraph("Item two", style="List Bullet")

    path = tmp_path / "sample.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def sample_docx_with_image(tmp_path: Path) -> Path:
    """Create a .docx containing an embedded image."""
    doc = Document()
    doc.add_heading("Document With Image", level=1)
    doc.add_paragraph("Below is an image:")

    # Create a small PNG in memory
    img_path = tmp_path / "test_image.png"
    _create_small_png(img_path)
    doc.add_picture(str(img_path), width=Inches(1))
    doc.add_paragraph("Above was the image.")

    path = tmp_path / "with_image.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def sample_docx_with_table(tmp_path: Path) -> Path:
    """Create a .docx containing a table."""
    doc = Document()
    doc.add_heading("Table Document", level=1)
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "100"
    table.cell(2, 0).text = "Beta"
    table.cell(2, 1).text = "200"

    path = tmp_path / "with_table.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def sample_pdf(tmp_path: Path) -> Path:
    """Create a simple PDF with headings and text."""
    doc = fitz.open()
    page = doc.new_page()
    # Title
    page.insert_text((72, 72), "Test PDF Document", fontsize=24)
    # Body
    page.insert_text((72, 120), "This is a test paragraph in the PDF.", fontsize=12)
    page.insert_text((72, 150), "Second line of text.", fontsize=12)

    path = tmp_path / "sample.pdf"
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def sample_pdf_with_image(tmp_path: Path) -> Path:
    """Create a PDF containing an embedded image."""
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "PDF With Image", fontsize=18)

    # Create a small PNG and insert it
    img_path = tmp_path / "test_image.png"
    _create_small_png(img_path)
    rect = fitz.Rect(72, 100, 172, 200)
    page.insert_image(rect, filename=str(img_path))

    path = tmp_path / "with_image.pdf"
    doc.save(str(path))
    doc.close()
    return path


def _create_small_png(path: Path) -> None:
    """Create a minimal 2x2 red PNG file."""
    import struct
    import zlib

    def _chunk(chunk_type: bytes, data: bytes) -> bytes:
        c = chunk_type + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)

    signature = b"\x89PNG\r\n\x1a\n"
    ihdr = _chunk(b"IHDR", struct.pack(">IIBBBBB", 2, 2, 8, 2, 0, 0, 0))
    # 2x2 RGB image, all red
    raw = b""
    for _ in range(2):
        raw += b"\x00"  # filter byte
        raw += b"\xff\x00\x00" * 2  # 2 red pixels
    idat = _chunk(b"IDAT", zlib.compress(raw))
    iend = _chunk(b"IEND", b"")
    path.write_bytes(signature + ihdr + idat + iend)
